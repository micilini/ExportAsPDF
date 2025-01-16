import json
import base64
import sys  # Import sys for command-line arguments
from io import BytesIO
import tempfile
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image
from docx.shared import RGBColor  # Import RGBColor
from docx.enum.text import WD_COLOR_INDEX  # Import color index for highlighting
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.oxml import parse_xml
from docx2pdf import convert
import os
import winreg
import shutil
import subprocess
from docx.shared import Inches


def is_microsoft_office_installed():
    try:
        key = winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\Microsoft\Office")
        return True
    except FileNotFoundError:
        return False

import os

def is_libreoffice_installed():
    paths_to_check = [
        r"C:/Program Files/LibreOffice/program/soffice.exe",
        r"/usr/bin/soffice",  # For Linux
        r"/usr/local/bin/soffice"  # Alternative Linux path
    ]
    return any(os.path.isfile(path) for path in paths_to_check)



def customize_run(run, font_size=None, color=None):
    """
    Customize a run with the given font size and color.
    :param run: The run object to customize.
    :param font_size: Font size in points.
    :param color: Color as an RGB tuple (e.g., (255, 0, 0) for red).
    """
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)

#when we add any hyper links this added twice one as a hyperlink and the another as normal text, so this function will delete duplicated normal text.
def clean_duplicate_links(doc):
    for paragraph in doc.paragraphs:
        # Set to store text of hyperlinks
        hyperlink_texts = set()
        runs_to_remove = []

        # Iterate through the paragraph's XML to find hyperlink elements
        for element in paragraph._element:
            if element.tag.endswith("hyperlink"):
                # Extract text inside the hyperlink
                for child in element:
                    if child.tag.endswith("r"):  # Check if it's a run element
                        for t in child:
                            if t.tag.endswith("t"):  # Text element
                                hyperlink_text = t.text.strip()
                                hyperlink_texts.add(hyperlink_text)

        # Check runs for duplicates
        for run in paragraph.runs:
            if run.text.strip() in hyperlink_texts and run._element.getparent().tag != qn("w:hyperlink"):
                # Mark duplicate plain-text runs for removal
                runs_to_remove.append(run)

        # Remove duplicate runs
        for run in runs_to_remove:
            paragraph._element.remove(run._element)

#the function which added hyperlink
def add_hyperlink(paragraph, text, url):
    # Create the hyperlink tag
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    # Create a run and set its properties for the hyperlink
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    # Set font properties for the hyperlink text
    run_pr = OxmlElement("w:rPr")
    run_font = OxmlElement("w:sz")
    run_font.set(qn('w:val'), '24')  # Adjust size
    run_pr.append(run_font)
    run_bold = OxmlElement("w:b")
    run_bold.set(qn('w:val'), '1')  # Bold
    run_pr.append(run_bold)
    run_underline = OxmlElement("w:u")
    run_underline.set(qn('w:val'), 'single')  # Underline
    run_pr.append(run_underline)
    run_color = OxmlElement("w:color")
    run_color.set(qn('w:val'), '0000FF')  # Blue color for hyperlink
    run_pr.append(run_color)
    run.append(run_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)

    # Add the hyperlink to the paragraph
    paragraph._element.append(hyperlink)



# Load JSON file
def load_json(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return json.load(file)

#this one for make quote like textbox
def style_as_textbox(paragraph, background_color="D9D9D9"):
    """
    Styles a paragraph to look like a textbox with borders and background shading.
    :param paragraph: The paragraph to style.
    :param background_color: Background color in hex (default is light gray 'D9D9D9').
    """
    # Get the paragraph's XML
    p = paragraph._element

    # Create a border element
    pPr = p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")

    for border_type in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{border_type}")
        border.set(qn("w:val"), "single")  # Border type
        border.set(qn("w:sz"), "4")  # Border width
        border.set(qn("w:space"), "1")  # Space between border and text
        border.set(qn("w:color"), "000000")  # Black border color
        borders.append(border)

    pPr.append(borders)

    # Add background shading
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), background_color)  # Background color
    pPr.append(shading)


def process_block(block, doc):
    block_type = block.get('type')
    data = block.get('data', {})

    if block_type == 'header':
        level = data.get('level', 1)
        text = data.get('text', '')
        paragraph = doc.add_heading(level=level)
        add_formatted_text(paragraph, text)
        for run in paragraph.runs:
            if level==1:
                customize_run(run, font_size=32, color=(0, 0, 0))  # Black headers, size decreases with level
            else:
                customize_run(run, font_size=20 + (6 - level) * 2, color=(0, 0, 0))  # Black headers, size decreases with level
        
    elif block_type == 'paragraph':
        paragraph = doc.add_paragraph()
        add_formatted_text(paragraph, data.get('text', ''))
        for run in paragraph.runs:
           run.font.size = Pt(18)

    elif block_type == 'list':
        style = data.get('style', 'unordered')
        for item in data.get('items', []):
            paragraph = doc.add_paragraph(style='List Bullet' if style == 'unordered' else 'List Number')
            add_formatted_text(paragraph, item)
            for run in paragraph.runs:
               run.font.size = Pt(18)

    elif block_type == 'checklist':
        for item in data.get('items', []):
            status = "✅" if item.get('checked') else "⬜"
            paragraph = doc.add_paragraph(f"{status} ")
            add_formatted_text(paragraph, item.get('text', ''))
            for run in paragraph.runs:
              run.font.size = Pt(18)

    elif block_type == 'quote':
        paragraph = doc.add_paragraph(style='Intense Quote')
        add_formatted_text(paragraph, f"“{data.get('text', '')}”")
        style_as_textbox(paragraph, background_color="F4F4F4")  # Light gray for quotes
        for run in paragraph.runs:
           run.font.size = Pt(18)
        caption = data.get('caption', '')
        if caption:
            caption_paragraph = doc.add_paragraph(style='Caption')
            add_formatted_text(caption_paragraph, f"- {caption}")
            for run in paragraph.runs:
              run.font.size = Pt(14)

    elif block_type == 'warning': 
        title = data.get('title', '')
        message = data.get('message', '')
        paragraph = doc.add_paragraph(style='Quote')
        add_formatted_text(paragraph, f"⚠️ {title}: {message}")
        style_as_textbox(paragraph, background_color="FFF2CC")  # Light yellow for warnings
        for run in paragraph.runs:
           run.font.size = Pt(18)

    elif block_type == 'code':
        paragraph = doc.add_paragraph(data.get('code', ''), style='Normal') 
        for run in paragraph.runs:
           run.font.size = Pt(16)

    elif block_type == 'delimiter':
        paragraph = doc.add_paragraph(style='Title')
        paragraph.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
        add_formatted_text(paragraph, '***')  # Visual separation

    elif block_type == 'table':
        table_data = data.get('content', [])
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        for i, row in enumerate(table_data):
            for j, cell in enumerate(row):
                cell_paragraph = table.cell(i, j).paragraphs[0]
                add_formatted_text(cell_paragraph, cell)
                for run in cell_paragraph.runs:
                  run.font.size = Pt(18)

    elif block_type == 'image':
        add_image_to_fit_page_or_original(data.get('url'), doc)


# Add formatted text with inline styles
def add_formatted_text(paragraph, text):
    from html.parser import HTMLParser

    class InlineStyleParser(HTMLParser):
        def __init__(self):
            super().__init__()
            self.current_tag = None
            self.current_attrs = {}

        def handle_starttag(self, tag, attrs):
            self.current_tag = tag
            self.current_attrs = dict(attrs)

        def handle_endtag(self, tag):
            self.current_tag = None

        def handle_data(self, data):
            run = paragraph.add_run(data)

            if self.current_tag == "i":
               run.italic = True
            elif self.current_tag == "b":
               run.bold = True
            elif self.current_tag == "u":
               run.underline = True
            elif self.current_tag == "code":
               run.font.color.rgb = RGBColor(255, 0, 0)  # Red for code
            elif self.current_tag == "mark":
               run.font.highlight_color = WD_COLOR_INDEX.YELLOW  # Yellow highlight for sublime
            elif self.current_tag == "a":
          # Check for href only if it's a new <a> tag to avoid processing the same link twice
             if "href" in self.current_attrs:
               add_hyperlink(paragraph, data, self.current_attrs["href"])
          # Reset to None after processing
               self.current_tag = None
            else:
             # Reset formatting for normal text
               run.font.color.rgb = None
 
    parser = InlineStyleParser()
    parser.feed(text)


# Add images to the document
def add_image_to_fit_page_or_original(base64_string, doc, page_width=6.0, page_height=8.0):
    """
    Add an image to the Word document, resizing only if it exceeds the page size.
    
    Args:
        base64_string (str): The base64-encoded image string.
        doc (Document): The Word document object.
        page_width (float): Maximum width of the image in inches (default is 6 inches for standard margins).
        page_height (float): Maximum height of the image in inches (default is 8 inches for standard margins).
    """
    if base64_string.startswith("data:image/"):
        # Decode the base64 image
        img_data = base64.b64decode(base64_string.split(",")[1])
        img = Image.open(BytesIO(img_data))
        
        # Get the original dimensions of the image
        original_width, original_height = img.size  # In pixels
        dpi = img.info.get("dpi", (96, 96))[0]  # Default DPI is 96 if not specified
        img_width_in_inches = original_width / dpi
        img_height_in_inches = original_height / dpi
        
        # Check if the image needs resizing
        if img_width_in_inches > page_width or img_height_in_inches > page_height:
            # Calculate the scaling factors to fit the image within the page
            width_scale = page_width / img_width_in_inches
            height_scale = page_height / img_height_in_inches
            scale = min(width_scale, height_scale)  # Use the smaller scaling factor
            
            # Calculate the new dimensions
            scaled_width = img_width_in_inches * scale
            scaled_height = img_height_in_inches * scale
        else:
            # Use the original dimensions
            scaled_width = img_width_in_inches
            scaled_height = img_height_in_inches
        
        # Save the image temporarily
        img.save("temp_image.png")
        
        # Add visual separation before the image
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.add_run('\n')  # Visual separation
        
        # Add the image to the document with calculated dimensions
        doc.add_picture("temp_image.png", width=Inches(scaled_width), height=Inches(scaled_height))
        
        # Add visual separation after the image
        paragraph = doc.add_paragraph()
        paragraph.add_run('\n')  # Visual separation

# Main function
def create_docx_from_json(json_data, output_docx_path):
    doc = Document()

    for block in json_data.get('blocks', []):
        process_block(block, doc)

    # Remove duplicate hyperlinks
    clean_duplicate_links(doc)

    # Save the DOCX file to BytesIO
    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    
    return doc_stream

#for microsoft word
def convert_docx_to_pdf_word(docx_stream):
    """
    Convert DOCX content from a BytesIO stream to PDF content as a BytesIO object.
    :param docx_stream: BytesIO stream containing DOCX data.
    :return: BytesIO object containing PDF data.
    """
    # Create temporary files for DOCX and PDF
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as docx_temp:
        docx_temp.write(docx_stream.getvalue())
        docx_temp_path = docx_temp.name

    pdf_stream = BytesIO()
    try:
        # Convert DOCX to PDF
        pdf_temp_path = f"{docx_temp_path[:-5]}.pdf"
        convert(docx_temp_path)

        # Read the resulting PDF into a BytesIO object
        with open(pdf_temp_path, "rb") as pdf_file:
            pdf_stream.write(pdf_file.read())

    finally:
        # Clean up temporary files
        if os.path.exists(docx_temp_path):
            os.remove(docx_temp_path)
        if os.path.exists(pdf_temp_path):
            os.remove(pdf_temp_path)

    pdf_stream.seek(0)  # Rewind the stream to the beginning
    return pdf_stream

#for libre office
def find_libreoffice():
    """
    Find the LibreOffice executable path from a list of potential paths.
    :return: The valid path to LibreOffice executable or None if not found.
    """
    paths = [
        r"C:/Program Files/LibreOffice/program/soffice.exe",  # Windows
        r"/usr/bin/soffice",  # Linux
        r"/usr/local/bin/soffice"  # Alternative Linux path
    ]
    
    for path in paths:
        if os.path.exists(path):
            return path
    
    # Fallback to default system PATH lookup
    return shutil.which("soffice")

def convert_docx_to_pdf_libre(docx_binary_data):
    """
    Convert DOCX content from binary data to PDF content as binary data.
    :param docx_binary_data: Binary content of the DOCX file.
    :return: BytesIO object containing PDF data.
    """
    # Create a temporary file for the DOCX binary data
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as docx_temp:
        docx_temp.write(docx_binary_data.read())
        docx_temp_path = docx_temp.name

    pdf_stream = BytesIO()
    try:
        # Define the output PDF path (temporary file)
        pdf_temp_path = f"{docx_temp_path[:-5]}.pdf"

        # Find LibreOffice path
        libreoffice_path = find_libreoffice()
        if not libreoffice_path:
            raise EnvironmentError("LibreOffice is not installed or not found in the specified paths.")

        # Convert using LibreOffice
        command = [
            libreoffice_path,
            "--headless",
            "--convert-to", "pdf",
            "--outdir", os.path.dirname(pdf_temp_path),
            docx_temp_path
        ]
        subprocess.run(command, check=True)

        # Read the resulting PDF into a BytesIO object
        with open(pdf_temp_path, "rb") as pdf_file:
            pdf_stream.write(pdf_file.read())

    except subprocess.CalledProcessError as e:
        print(f"Error during conversion: {e}")
    finally:
        # Clean up temporary DOCX and PDF files
        if os.path.exists(docx_temp_path):
            os.remove(docx_temp_path)
        if os.path.exists(pdf_temp_path):
            os.remove(pdf_temp_path)

    pdf_stream.seek(0)  # Rewind the stream to the beginning
    return pdf_stream


def detect_office_suite(docx_stream):
    try:
        if is_microsoft_office_installed():
            return convert_docx_to_pdf_word(docx_stream)
    except Exception as e:
        print(f"Error with Microsoft Office: {e}")
        
    try:
        if is_libreoffice_installed():
            return convert_docx_to_pdf_libre(docx_stream)
    except Exception as e:
        print(f"Error with LibreOffice: {e}")


if __name__ == "__main__":
    json_file_path = sys.argv[1]
    pdf_output_path = "output.pdf"

    # Load JSON data
    with open(json_file_path, 'r', encoding='utf-8') as file:
        json_data = json.load(file)

    # Create DOCX from JSON
    docx_stream = create_docx_from_json(json_data, "output.docx")
    pdf=detect_office_suite(docx_stream)


    sys.stdout.buffer.write(pdf.getvalue())
#for test
""" def save_pdf_stream_to_file(pdf_stream: bytes, output_path: str):
    with open(output_path, "wb") as pdf_file:
        pdf_file.write(pdf_stream)
    print(f"PDF saved successfully to {output_path}")

save_pdf_stream_to_file(pdf.getvalue(),"test1.pdf") """